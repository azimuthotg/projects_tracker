#!/usr/bin/env python3
# build_manual.py - สร้างคู่มือการใช้งาน ระบบติดตามโครงการ (Word .docx)
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENSHOTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'คู่มือการใช้งานระบบติดตามโครงการ.docx')

def set_page_a4(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

def set_thai(run, size=16, bold=False, color=None):
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h(doc, text, level=1):
    sizes  = {1: 22, 2: 18, 3: 16}
    colors = {1: (30,64,175), 2: (30,64,175), 3: (55,65,81)}
    before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_thai(r, size=sizes[level], bold=True, color=colors[level])
    return p

def para(doc, text, size=15, indent=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    set_thai(r, size=size, bold=bold, color=color)
    return p

def step(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_thai(r, size=15)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_thai(r, size=15)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run('\U0001f4cc ' + text)
    set_thai(r, size=14, color=(120,53,15))
    return p

def img(doc, filename, width_cm=14, caption=None):
    path = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(path):
        para(doc, '[ภาพ: ' + filename + ' ไม่พบ]', size=12, color=(200,0,0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run('ภาพ: ' + caption)
        set_thai(r, size=13, color=(107,114,128))

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def role_table(doc):
    rows = [
        ('เจ้าหน้าที่ (Staff)',       'ดูโครงการที่ตนเองรับผิดชอบ, บันทึกเบิกจ่าย'),
        ('เจ้าหน้าที่แผน (Planner)', 'ดูโครงการทุกโครงการในแผนก, สร้าง/แก้ไขโครงการ, บันทึกเบิกจ่าย'),
        ('หัวหน้างาน (Head)',         'ทุกอย่างของ Planner + อนุมัติ/ปฏิเสธการเบิกจ่าย'),
        ('ผู้ดูแลระบบ (Admin)',       'ทุกอย่างของ Head + จัดการผู้ใช้, แผนก, ปีงบประมาณ'),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, txt in enumerate(['บทบาท', 'สิทธิ์การใช้งาน']):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(txt)
        set_thai(r, size=15, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, '1E40AF')
    for i, (role, perm) in enumerate(rows):
        cells = table.rows[i+1].cells
        r1 = cells[0].paragraphs[0].add_run(role)
        set_thai(r1, size=15, bold=True)
        r2 = cells[1].paragraphs[0].add_run(perm)
        set_thai(r2, size=15)
        if i % 2 == 0:
            shade_cell(cells[0], 'EFF6FF')
            shade_cell(cells[1], 'EFF6FF')
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)
    doc.add_paragraph()

def status_table(doc):
    rows = [
        ('ร่าง (draft)',     'โครงการที่ยังไม่พร้อมใช้งาน'),
        ('ยังไม่ดำเนินการ', 'อนุมัติแล้ว รอเริ่มดำเนินการ'),
        ('ดำเนินการ',        'กำลังดำเนินการอยู่'),
        ('เสร็จสิ้น',        'ดำเนินการครบถ้วนแล้ว'),
        ('ยกเลิก',           'ยกเลิกโครงการ'),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    for i, txt in enumerate(['สถานะ', 'ความหมาย']):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(txt)
        set_thai(r, size=15, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, '1E40AF')
    for i, (st, desc) in enumerate(rows):
        cells = table.rows[i+1].cells
        set_thai(cells[0].paragraphs[0].add_run(st), size=15, bold=True)
        set_thai(cells[1].paragraphs[0].add_run(desc), size=15)
    doc.add_paragraph()

# ─── cover ───────────────────────────────────────────────────────────────────

def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('คู่มือการใช้งาน'), size=32, bold=True, color=(30,64,175))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('ระบบติดตามแผนงานและงบประมาณ'), size=26, bold=True, color=(30,64,175))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('Project Tracker'), size=20, color=(107,114,128))
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('สำนักวิทยบริการ มหาวิทยาลัยนครพนม'), size=18, bold=True, color=(55,65,81))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('Academic Resource Center, Nakhon Phanom University'), size=15, color=(107,114,128))
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('เวอร์ชัน 1.0  |  มีนาคม 2569'), size=14, color=(107,114,128))
    doc.add_page_break()

def ch1(doc):
    h(doc, 'บทที่ 1  บทนำ', 1)
    h(doc, '1.1  วัตถุประสงค์ของระบบ', 2)
    para(doc, 'ระบบติดตามแผนงานและงบประมาณ (Project Tracker) พัฒนาขึ้นเพื่อแก้ไขปัญหาการจัดการข้อมูลโครงการและงบประมาณของสำนักวิทยบริการ มหาวิทยาลัยนครพนม ซึ่งเดิมพึ่งพาเอกสารกระดาษและระบบ ERP ที่เข้าถึงได้เฉพาะเจ้าหน้าที่แผนเท่านั้น')
    para(doc, 'ระบบนี้ช่วยให้บุคลากรทุกคนสามารถ:')
    bullet(doc, 'ดูข้อมูลโครงการและงบประมาณที่รับผิดชอบได้ตลอดเวลา')
    bullet(doc, 'ติดตามยอดคงเหลืองบประมาณแต่ละกิจกรรมแบบ real-time')
    bullet(doc, 'บันทึกการเบิกจ่ายและส่งขออนุมัติออนไลน์')
    bullet(doc, 'รับการแจ้งเตือนผ่าน LINE เมื่องบใกล้เต็มหรือใกล้ถึงกำหนด')
    bullet(doc, 'ออกรายงานงบประมาณในรูปแบบ Excel และ PDF')
    h(doc, '1.2  กลุ่มผู้ใช้งานและบทบาท', 2)
    para(doc, 'ระบบแบ่งผู้ใช้งานออกเป็น 4 บทบาท แต่ละบทบาทมีสิทธิ์การเข้าถึงแตกต่างกัน:')
    role_table(doc)
    note(doc, 'บทบาทของผู้ใช้งานถูกกำหนดโดยผู้ดูแลระบบ หากต้องการเปลี่ยนแปลงบทบาท ให้ติดต่อผู้ดูแลระบบ')
    h(doc, '1.3  ความต้องการของระบบ', 2)
    bullet(doc, 'เบราว์เซอร์: Google Chrome, Microsoft Edge หรือ Firefox เวอร์ชันล่าสุด')
    bullet(doc, 'การเชื่อมต่ออินเทอร์เน็ต: ภายในเครือข่ายมหาวิทยาลัยหรือ VPN')
    bullet(doc, 'URL: https://lib.npu.ac.th/projects/')
    bullet(doc, 'บัญชีผู้ใช้: ใช้ชื่อผู้ใช้และรหัสผ่านเดียวกับระบบมหาวิทยาลัย (NPU)')
    doc.add_page_break()

def ch2(doc):
    h(doc, 'บทที่ 2  การเข้าสู่ระบบและออกจากระบบ', 1)
    h(doc, '2.1  การเข้าสู่ระบบ', 2)
    para(doc, 'เปิดเบราว์เซอร์แล้วไปที่ https://lib.npu.ac.th/projects/ จะพบหน้าเข้าสู่ระบบดังนี้')
    img(doc, '01_login.png', width_cm=12, caption='หน้าเข้าสู่ระบบ')
    para(doc, 'ขั้นตอนการเข้าสู่ระบบ:')
    step(doc, 'กรอกชื่อผู้ใช้ ด้วยเลขบัตรประชาชน หรือชื่อผู้ใช้งาน NPU ของท่าน')
    step(doc, 'กรอกรหัสผ่าน ซึ่งเป็นรหัสผ่านเดียวกับที่ใช้ล็อกอินระบบมหาวิทยาลัย')
    step(doc, 'คลิกปุ่ม "เข้าสู่ระบบ"')
    step(doc, 'ระบบจะนำไปยังหน้าแดชบอร์ดโดยอัตโนมัติ')
    note(doc, 'หากเข้าสู่ระบบครั้งแรกและยังไม่มีบัญชี ระบบจะสร้างบัญชีให้อัตโนมัติโดยดึงข้อมูลจาก NPU Directory ผู้ดูแลระบบจะต้องอนุมัติบัญชีก่อนจึงจะใช้งานได้')
    h(doc, '2.2  การออกจากระบบ', 2)
    para(doc, 'คลิกปุ่ม "ออกจากระบบ" ที่มุมบนขวาของทุกหน้า เพื่อออกจากระบบอย่างปลอดภัย')
    note(doc, 'ควรออกจากระบบทุกครั้งหลังใช้งานเสร็จ โดยเฉพาะเมื่อใช้คอมพิวเตอร์สาธารณะ')
    doc.add_page_break()

def ch3(doc):
    h(doc, 'บทที่ 3  แดชบอร์ดและการนำทาง', 1)
    h(doc, '3.1  หน้าแดชบอร์ด', 2)
    para(doc, 'เมื่อเข้าสู่ระบบสำเร็จ ระบบจะแสดงหน้าแดชบอร์ดซึ่งสรุปภาพรวมของโครงการและงบประมาณ')
    img(doc, '02_dashboard.png', width_cm=15, caption='หน้าแดชบอร์ดหลัก')
    para(doc, 'ส่วนประกอบของแดชบอร์ด:')
    bullet(doc, 'การ์ด 4 ใบด้านบน: แสดงจำนวนโครงการทั้งหมด, โครงการที่กำลังดำเนินการ, งบประมาณรวม, และรายการรออนุมัติ')
    bullet(doc, 'แถบความคืบหน้างบประมาณ: แสดงยอดใช้จ่ายรวมของปีงบประมาณปัจจุบัน')
    bullet(doc, 'รายการที่ต้องติดตาม: แจ้งเตือนกิจกรรมที่เลย deadline, งบเต็ม, ไม่มีรายงาน และถึงเวลาแต่ยังไม่เริ่ม')
    h(doc, '3.2  แถบเมนูด้านซ้าย (Sidebar)', 2)
    para(doc, 'เมนูทางด้านซ้ายมือใช้สำหรับนำทางไปยังส่วนต่าง ๆ ของระบบ:')
    bullet(doc, 'แดชบอร์ด — ภาพรวมโครงการและงบประมาณ')
    bullet(doc, 'โครงการ — รายการโครงการทั้งหมดที่มีสิทธิ์เข้าถึง')
    bullet(doc, 'ปฏิทินโครงการ — แสดงกำหนดการในรูปแบบ Gantt Chart')
    bullet(doc, 'งบประมาณ — รายการเบิกจ่ายทั้งหมด')
    bullet(doc, 'อนุมัติรายการ — (หัวหน้า/Admin) รายการรออนุมัติ')
    bullet(doc, 'ภาพรวมผู้บริหาร — (Planner/Head/Admin) สรุปภาพรวมทุกแผนก')
    bullet(doc, 'งานของฉัน — กิจกรรมและการเบิกจ่ายที่เกี่ยวข้องกับตนเอง')
    bullet(doc, 'รายงาน — ออกรายงานงบประมาณและค่าใช้จ่าย')
    note(doc, 'เมนูที่แสดงจะขึ้นอยู่กับบทบาทของผู้ใช้งาน เจ้าหน้าที่ทั่วไปจะเห็นเมนูน้อยกว่าหัวหน้างานหรือผู้ดูแลระบบ')
    doc.add_page_break()

def ch4(doc):
    h(doc, 'บทที่ 4  การจัดการโครงการ', 1)
    h(doc, '4.1  รายการโครงการ', 2)
    para(doc, 'คลิกเมนู "โครงการ" เพื่อดูรายการโครงการทั้งหมดที่มีสิทธิ์เข้าถึง')
    img(doc, '03_project_list.png', width_cm=15, caption='หน้ารายการโครงการ')
    para(doc, 'การค้นหาและกรองข้อมูล:')
    bullet(doc, 'เลือกปีงบประมาณ เพื่อกรองโครงการตามปี')
    bullet(doc, 'เลือกสถานะ เพื่อดูโครงการตามสถานะ')
    bullet(doc, 'พิมพ์ชื่อหรือรหัสโครงการในช่องค้นหา แล้วคลิก "ค้นหา"')
    bullet(doc, 'คลิก "ล้าง" เพื่อยกเลิกการกรอง')
    para(doc, 'ข้อมูลในตาราง:')
    bullet(doc, 'รหัส/ลำดับ: รหัสโครงการ')
    bullet(doc, 'ชื่อโครงการ: ชื่อ, แผนก, ผู้รับผิดชอบ, ระยะเวลา')
    bullet(doc, 'แหล่งเงิน: งบแยกตามเงินแผ่นดิน, เงินสะสม, เงินรายได้')
    bullet(doc, 'งบประมาณ: งบรวม, ยอดใช้ไป/คงเหลือ และเปอร์เซ็นต์การใช้')
    bullet(doc, 'สถานะ: สีแสดงสถานะปัจจุบันของโครงการ')
    bullet(doc, 'คลิกปุ่ม > หน้าชื่อโครงการ เพื่อขยายดูรายการกิจกรรม')
    h(doc, '4.2  รายละเอียดโครงการ', 2)
    para(doc, 'คลิกชื่อโครงการในรายการ เพื่อดูรายละเอียดโครงการ')
    img(doc, '04_project_detail.png', width_cm=15, caption='หน้ารายละเอียดโครงการ')
    para(doc, 'หน้านี้แสดงข้อมูล:')
    bullet(doc, 'สรุปงบประมาณ: งบรวม, จัดสรรแล้ว, ใช้ไป, คงเหลือ — แยกตามแหล่งเงิน')
    bullet(doc, 'แถบความคืบหน้า: แสดงเปอร์เซ็นต์การใช้งบประมาณรวม')
    bullet(doc, 'รายละเอียดโครงการ: คำอธิบาย, ผู้รับผิดชอบ, เอกสารแนบ')
    bullet(doc, 'ตารางกิจกรรม: รายการกิจกรรมทั้งหมดพร้อมงบประมาณและสถานะ')
    bullet(doc, 'รายการเบิกจ่ายล่าสุด: แสดงการเบิกจ่าย 5 รายการล่าสุด')
    h(doc, '4.3  สถานะโครงการ', 2)
    status_table(doc)
    h(doc, '4.4  การสร้างโครงการใหม่ (Planner/Head/Admin)', 2)
    para(doc, 'ผู้ใช้งานที่มีบทบาท Planner ขึ้นไปสามารถสร้างโครงการใหม่ได้:')
    step(doc, 'คลิกปุ่ม "+ สร้างโครงการใหม่" มุมบนขวาของหน้ารายการโครงการ')
    step(doc, 'กรอกข้อมูลโครงการ: รหัส, ชื่อ, คำอธิบาย, วันเริ่ม-สิ้นสุด')
    step(doc, 'ระบุงบประมาณแยกตามแหล่งเงิน (เงินแผ่นดิน / เงินสะสม / เงินรายได้)')
    step(doc, 'เลือกผู้รับผิดชอบและผู้รับการแจ้งเตือน')
    step(doc, 'คลิก "บันทึก" เพื่อสร้างโครงการ')
    note(doc, 'ผลรวมงบประมาณกิจกรรมทั้งหมดต้องไม่เกินงบประมาณรวมของโครงการ')
    doc.add_page_break()

def ch5(doc):
    h(doc, 'บทที่ 5  การจัดการกิจกรรม', 1)
    h(doc, '5.1  รายละเอียดกิจกรรม', 2)
    para(doc, 'คลิกชื่อกิจกรรมในหน้ารายละเอียดโครงการ เพื่อดูข้อมูลกิจกรรมนั้น ๆ')
    img(doc, '05_activity_detail.png', width_cm=15, caption='หน้ารายละเอียดกิจกรรม')
    para(doc, 'หน้ารายละเอียดกิจกรรมประกอบด้วย:')
    bullet(doc, 'สรุปงบประมาณ: จัดสรร, ใช้ไป, คงเหลือ — แยกตามแหล่งเงิน')
    bullet(doc, 'แถบความคืบหน้า: สีเปลี่ยนตามเปอร์เซ็นต์ (เขียว <70%, เหลือง <90%, แดง >=90%)')
    bullet(doc, 'ตารางรายการเบิกจ่าย: แสดงค่าใช้จ่ายทั้งหมดพร้อมสถานะ')
    bullet(doc, 'ไฟล์แนบ: สามารถอัปโหลดเอกสารหลักฐานการจ่ายเงิน')
    h(doc, '5.2  การเปลี่ยนสถานะกิจกรรม (Planner/Head/Admin)', 2)
    bullet(doc, 'คลิกปุ่ม "เปิดดำเนินการ" — เปลี่ยนเป็นสถานะ ดำเนินการ')
    bullet(doc, 'คลิกปุ่ม "เสร็จสิ้นกิจกรรม" — เปลี่ยนเป็นสถานะ เสร็จสิ้น')
    bullet(doc, 'คลิกปุ่ม "เปิดใหม่" — เปิดกิจกรรมที่เสร็จสิ้นแล้วกลับมาดำเนินการใหม่')
    h(doc, '5.3  การเพิ่มกิจกรรมใหม่ (Planner/Head/Admin)', 2)
    step(doc, 'เปิดหน้ารายละเอียดโครงการ')
    step(doc, 'คลิกปุ่ม "+ เพิ่มกิจกรรม" ในส่วนกิจกรรม')
    step(doc, 'กรอกข้อมูล: ชื่อกิจกรรม, ลำดับที่, วันเริ่ม-สิ้นสุด, งบประมาณที่จัดสรร')
    step(doc, 'เลือกแหล่งเงิน (เงินแผ่นดิน / เงินสะสม / เงินรายได้) หรือทำเครื่องหมาย "ไม่ใช้งบประมาณ"')
    step(doc, 'คลิก "บันทึก"')
    note(doc, 'ผลรวมงบประมาณกิจกรรมทั้งหมดต้องไม่เกินงบประมาณรวมของโครงการ')
    doc.add_page_break()

def ch6(doc):
    h(doc, 'บทที่ 6  การบันทึกการเบิกจ่าย', 1)
    h(doc, '6.1  รายการเบิกจ่ายทั้งหมด', 2)
    para(doc, 'คลิกเมนู "งบประมาณ" เพื่อดูรายการเบิกจ่ายของโครงการที่ท่านมีสิทธิ์')
    img(doc, '07_expense_list.png', width_cm=15, caption='หน้ารายการเบิกจ่าย')
    para(doc, 'สามารถกรองรายการได้ตามโครงการ, กิจกรรม, สถานะ และวันที่')
    h(doc, '6.2  การบันทึกรายการเบิกจ่ายใหม่', 2)
    para(doc, 'มี 2 วิธีในการบันทึกการเบิกจ่าย:')
    bullet(doc, 'วิธีที่ 1: จากหน้ากิจกรรม — คลิกปุ่ม "บันทึกเบิกจ่าย" สีน้ำเงินมุมบนขวา')
    bullet(doc, 'วิธีที่ 2: จากเมนูงบประมาณ — คลิกปุ่ม "+ บันทึกการเบิกจ่าย"')
    img(doc, '06_expense_form.png', width_cm=14, caption='แบบฟอร์มบันทึกการเบิกจ่าย')
    para(doc, 'กรอกข้อมูลในแบบฟอร์ม:')
    step(doc, 'เลือกโครงการและกิจกรรมที่ต้องการเบิก')
    step(doc, 'กรอกรายละเอียดของค่าใช้จ่าย')
    step(doc, 'กรอกจำนวนเงิน (ต้องไม่เกินงบคงเหลือของกิจกรรม)')
    step(doc, 'เลือกแหล่งเงินที่ใช้เบิก')
    step(doc, 'กรอกวันที่เบิกจ่าย และเลขที่ใบเสร็จ (ถ้ามี)')
    step(doc, 'แนบไฟล์หลักฐาน (PDF หรือรูปภาพ) ถ้ามี')
    step(doc, 'คลิก "บันทึก"')
    note(doc, 'เจ้าหน้าที่แผน (Planner) และผู้ดูแลระบบ (Admin): รายการจะได้รับการอนุมัติอัตโนมัติทันที\nเจ้าหน้าที่ (Staff): รายการจะมีสถานะ "รออนุมัติ" จนกว่าหัวหน้าจะอนุมัติ')
    h(doc, '6.3  สถานะการเบิกจ่าย', 2)
    bullet(doc, 'รออนุมัติ — บันทึกแล้ว รอหัวหน้าตรวจสอบ')
    bullet(doc, 'อนุมัติ — ได้รับการอนุมัติแล้ว ยอดเงินถูกหักจากงบประมาณ')
    bullet(doc, 'ปฏิเสธ — ไม่ผ่านการอนุมัติ ยอดเงินไม่ถูกหัก')
    doc.add_page_break()

def ch7(doc):
    h(doc, 'บทที่ 7  การอนุมัติการเบิกจ่าย (หัวหน้างาน / Admin)', 1)
    h(doc, '7.1  รายการรออนุมัติ', 2)
    para(doc, 'หัวหน้างานและผู้ดูแลระบบจะเห็นเมนู "อนุมัติรายการ" ในแถบเมนูซ้าย ตัวเลขสีเหลืองแสดงจำนวนรายการที่รออนุมัติ')
    img(doc, '08_approval_list.png', width_cm=15, caption='หน้ารายการรออนุมัติ')
    h(doc, '7.2  การอนุมัติหรือปฏิเสธรายการ', 2)
    step(doc, 'คลิกเมนู "อนุมัติรายการ"')
    step(doc, 'คลิกชื่อรายการที่ต้องการตรวจสอบ')
    step(doc, 'ตรวจสอบรายละเอียด: จำนวนเงิน, วันที่, ใบเสร็จ, ไฟล์แนบ')
    step(doc, 'คลิก "อนุมัติ" หรือ "ปฏิเสธ"')
    step(doc, 'หากปฏิเสธ กรอกเหตุผลในช่องหมายเหตุ')
    step(doc, 'คลิก "ยืนยัน"')
    note(doc, 'เมื่ออนุมัติแล้ว ระบบจะส่งการแจ้งเตือนผ่าน LINE ไปยังผู้บันทึกรายการนั้นโดยอัตโนมัติ และหากยอดใช้จ่ายถึงเกณฑ์ที่ตั้งไว้ จะมีการแจ้งเตือนงบประมาณใกล้เต็มด้วย')
    doc.add_page_break()

def ch8(doc):
    h(doc, 'บทที่ 8  งานของฉัน', 1)
    para(doc, 'หน้า "งานของฉัน" รวบรวมสิ่งที่เกี่ยวข้องกับตัวท่านไว้ในที่เดียว')
    img(doc, '09_my_tasks.png', width_cm=15, caption='หน้างานของฉัน')
    para(doc, 'แบ่งออกเป็น 4 ส่วน:')
    bullet(doc, 'กิจกรรมเลย deadline — กิจกรรมที่เลยกำหนดเสร็จแล้วแต่ยังไม่ปิด')
    bullet(doc, 'กิจกรรมใกล้ถึงกำหนด — กิจกรรมที่จะครบกำหนดใน 7 วัน')
    bullet(doc, 'กิจกรรมที่รับผิดชอบ — กิจกรรม active ทั้งหมดที่ท่านเป็นผู้รับผิดชอบ')
    bullet(doc, 'รายการเบิกจ่ายของฉัน — รายการที่ท่านบันทึกและยังรออนุมัติ')
    note(doc, 'หัวหน้างานและ Admin จะมีส่วน "คิวอนุมัติ" เพิ่มเติม แสดงรายการทั้งหมดที่รอการอนุมัติจากตน')
    doc.add_page_break()

def ch9(doc):
    h(doc, 'บทที่ 9  ปฏิทินโครงการ (Gantt Chart)', 1)
    para(doc, 'คลิกเมนู "ปฏิทินโครงการ" เพื่อดูกำหนดการโครงการในรูปแบบ Gantt Chart')
    img(doc, '10_timeline.png', width_cm=15, caption='ปฏิทินโครงการแบบ Gantt Chart')
    para(doc, 'วิธีใช้ปฏิทินโครงการ:')
    bullet(doc, 'แถบแนวนอนแสดงระยะเวลาของแต่ละโครงการตามเดือน')
    bullet(doc, 'สีของแถบแสดงสถานะ: เขียว=ดำเนินการ, น้ำเงิน=ยังไม่เริ่ม, เทา=เสร็จสิ้น, แดง=ยกเลิก')
    bullet(doc, 'คลิกที่แถบเพื่อดูรายละเอียดโครงการ')
    bullet(doc, 'เลื่อนแนวนอนเพื่อดูเดือนที่ต้องการ')
    doc.add_page_break()

def ch10(doc):
    h(doc, 'บทที่ 10  รายงาน', 1)
    h(doc, '10.1  รายงานงบประมาณ', 2)
    para(doc, 'คลิกเมนู "รายงาน" เพื่อเข้าหน้ารายงานงบประมาณ')
    img(doc, '11_budget_report.png', width_cm=15, caption='หน้ารายงานงบประมาณ')
    para(doc, 'สามารถกรองรายงานตาม: ปีงบประมาณ, แผนก, สถานะโครงการ')
    para(doc, 'ตัวเลือกการส่งออก:')
    bullet(doc, 'คลิก "พิมพ์/PDF" — เปิดหน้าพิมพ์สำหรับพิมพ์หรือบันทึก PDF')
    bullet(doc, 'คลิก "ดาวน์โหลด Excel" — ดาวน์โหลดไฟล์ .xlsx')
    bullet(doc, 'คลิก "ดาวน์โหลด PDF" — ดาวน์โหลดไฟล์ .pdf')
    h(doc, '10.2  รายงานค่าใช้จ่าย', 2)
    para(doc, 'ในหน้ารายงาน คลิกแท็บ "รายงานค่าใช้จ่าย" เพื่อดูรายละเอียดการเบิกจ่ายแต่ละรายการ กรองได้ตามช่วงวันที่ โครงการ และสถานะ')
    h(doc, '10.3  รายงานโครงการรายโครงการ', 2)
    para(doc, 'ในหน้ารายละเอียดโครงการ คลิกปุ่ม "รายงาน" เพื่อดูรายงานของโครงการนั้น ๆ สามารถพิมพ์หรือดาวน์โหลด PDF ได้')
    h(doc, '10.4  ภาพรวมผู้บริหาร (Planner/Head/Admin)', 2)
    img(doc, '12_executive.png', width_cm=15, caption='หน้าภาพรวมผู้บริหาร')
    para(doc, 'หน้านี้แสดงสรุปภาพรวมของทุกแผนก ประกอบด้วยกราฟและตารางสรุปงบประมาณตามปีงบประมาณ แผนก และสถานะ')
    doc.add_page_break()

def ch11(doc):
    h(doc, 'บทที่ 11  การโอนงบประมาณระหว่างกิจกรรม (Planner/Head/Admin)', 1)
    para(doc, 'เมื่อกิจกรรมหนึ่งมีงบเหลือ และอีกกิจกรรมหนึ่งต้องการงบเพิ่ม สามารถโอนงบระหว่างกิจกรรมได้')
    img(doc, '14_budget_transfer.png', width_cm=14, caption='แบบฟอร์มโอนงบประมาณ')
    step(doc, 'เปิดหน้ารายละเอียดโครงการ')
    step(doc, 'คลิกปุ่ม "โอนงบ" (สีม่วง)')
    step(doc, 'เลือกกิจกรรมต้นทาง (กิจกรรมที่มีงบเหลือ)')
    step(doc, 'เลือกแหล่งเงินที่ต้องการโอน')
    step(doc, 'เลือกกิจกรรมปลายทาง (กิจกรรมที่รับงบ)')
    step(doc, 'กรอกจำนวนเงินที่ต้องการโอน')
    step(doc, 'กรอกเหตุผลประกอบการโอน')
    step(doc, 'คลิก "ยืนยันการโอนงบ"')
    note(doc, 'การโอนงบทุกครั้งจะถูกบันทึกใน Audit Log โดยอัตโนมัติ จำนวนเงินที่โอนต้องไม่เกินงบคงเหลือของกิจกรรมต้นทาง')
    doc.add_page_break()

def ch12(doc):
    h(doc, 'บทที่ 12  ข้อมูลส่วนตัวและการแจ้งเตือน LINE', 1)
    h(doc, '12.1  การดูข้อมูลส่วนตัว', 2)
    para(doc, 'คลิกชื่อผู้ใช้มุมบนขวา เพื่อเข้าหน้าข้อมูลส่วนตัว')
    img(doc, '13_profile.png', width_cm=14, caption='หน้าข้อมูลส่วนตัว')
    h(doc, '12.2  การตั้งค่าการแจ้งเตือน LINE', 2)
    para(doc, 'ในหน้าโปรไฟล์ ส่วนการแจ้งเตือน LINE สามารถตั้งค่าได้:')
    bullet(doc, 'เปิด/ปิดการแจ้งเตือนเมื่องบประมาณใกล้เต็ม')
    bullet(doc, 'เปิด/ปิดการแจ้งเตือนเมื่อใกล้ถึง deadline')
    bullet(doc, 'กำหนดเกณฑ์แจ้งเตือนงบ (ค่าเริ่มต้น 80%) — แจ้งเตือนเมื่อใช้งบถึงเปอร์เซ็นต์นี้')
    note(doc, 'ต้องให้ผู้ดูแลระบบเชื่อมต่อ LINE User ID ให้ก่อน จึงจะได้รับการแจ้งเตือนผ่าน LINE ได้ หากยังไม่มี LINE User ID หน้าโปรไฟล์จะแสดงว่า "ยังไม่ได้เชื่อมต่อ LINE"')
    doc.add_page_break()

def ch13(doc):
    h(doc, 'บทที่ 13  คำถามที่พบบ่อย (FAQ)', 1)
    faqs = [
        ('เข้าสู่ระบบไม่ได้ ขึ้นว่าชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง',
         'ตรวจสอบว่าใช้รหัสผ่านเดียวกับระบบมหาวิทยาลัย (NPU) หากลืมรหัสผ่านให้ติดต่อ IT ของมหาวิทยาลัย'),
        ('เข้าสู่ระบบแล้ว แต่บัญชีแสดงว่ารออนุมัติ',
         'บัญชีใหม่ต้องรอผู้ดูแลระบบอนุมัติก่อน ให้แจ้งผู้ดูแลระบบเพื่อเร่งการอนุมัติ'),
        ('ไม่เห็นโครงการที่ตนเองรับผิดชอบ',
         'ตรวจสอบว่าชื่อของท่านถูกเพิ่มเป็นผู้รับผิดชอบของโครงการ/กิจกรรมนั้นหรือไม่ หากไม่ใช่ให้แจ้ง Planner หรือ Admin'),
        ('บันทึกการเบิกจ่ายแล้ว แต่ยอดงบไม่ลด',
         'ยอดงบจะลดลงเฉพาะรายการที่อนุมัติแล้วเท่านั้น หากสถานะยังเป็น "รออนุมัติ" ยอดงบจะยังไม่เปลี่ยนแปลง'),
        ('ไม่ได้รับการแจ้งเตือน LINE',
         'ตรวจสอบว่า LINE User ID ถูกบันทึกในระบบแล้วหรือยัง และตรวจสอบการตั้งค่าการแจ้งเตือนในหน้าโปรไฟล์'),
        ('ต้องการลบโครงการ',
         'ไม่สามารถลบโครงการโดยตรงได้ ต้องส่งคำขอลบโครงการผ่านระบบ และรอ Admin อนุมัติ'),
        ('จำนวนเงินเบิกเกินงบคงเหลือ ระบบไม่ยอมบันทึก',
         'ระบบจะไม่อนุญาตให้บันทึกรายการที่เกินงบคงเหลือ หากต้องการงบเพิ่ม ให้ Planner/Head ทำการโอนงบมาก่อน'),
    ]
    for q, a in faqs:
        para(doc, 'Q: ' + q, size=15, bold=True, color=(30,64,175))
        para(doc, 'A: ' + a, size=15, indent=True)
        doc.add_paragraph()
    doc.add_page_break()

def ch14(doc):
    h(doc, 'บทที่ 14  การติดต่อผู้ดูแลระบบ', 1)
    para(doc, 'หากพบปัญหาในการใช้งานระบบ หรือต้องการความช่วยเหลือ ติดต่อได้ที่:')
    bullet(doc, 'หน่วยงาน: สำนักวิทยบริการ มหาวิทยาลัยนครพนม')
    bullet(doc, 'ที่อยู่เว็บไซต์: https://lib.npu.ac.th/projects/')
    para(doc, 'ผู้ดูแลระบบสามารถช่วยเหลือในเรื่อง:')
    bullet(doc, 'การตั้งค่าบัญชีและสิทธิ์การใช้งาน')
    bullet(doc, 'การเชื่อมต่อ LINE สำหรับการแจ้งเตือน')
    bullet(doc, 'การสร้างปีงบประมาณใหม่')
    bullet(doc, 'การจัดการแผนกและข้อมูลพื้นฐาน')
    bullet(doc, 'การอนุมัติบัญชีผู้ใช้งานใหม่')

def main():
    doc = Document()
    set_page_a4(doc)
    doc.styles['Normal'].font.name = 'TH Sarabun New'
    doc.styles['Normal'].font.size = Pt(15)
    cover(doc)
    ch1(doc); ch2(doc); ch3(doc); ch4(doc); ch5(doc)
    ch6(doc); ch7(doc); ch8(doc); ch9(doc); ch10(doc)
    ch11(doc); ch12(doc); ch13(doc); ch14(doc)
    doc.save(OUTPUT)
    print('บันทึกไฟล์: ' + OUTPUT)

if __name__ == '__main__':
    main()
