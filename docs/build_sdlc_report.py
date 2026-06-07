#!/usr/bin/env python3
# build_sdlc_report.py - สร้างรายงานการพัฒนาระบบ (SDLC) สำหรับประกอบการประเมิน
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENSHOTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'รายงานการพัฒนาระบบ_SDLC.docx')

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_page_a4(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

def set_thai(run, size=16, bold=False, color=None):
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h(doc, text, level=1):
    sizes  = {1: 22, 2: 18, 3: 16}
    colors = {1: (30,64,175), 2: (30,64,175), 3: (55,65,81)}
    before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_thai(r, size=sizes[level], bold=True, color=colors[level])
    return p

def para(doc, text, size=15, indent=False, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_thai(r, size=size, bold=bold, color=color)
    return p

def bullet(doc, text, size=15):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_thai(r, size=size)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def img(doc, filename, width_cm=14, caption=None):
    path = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(path):
        para(doc, f'[ภาพ: {filename} ไม่พบ]', size=12, color=(200,0,0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run('ภาพที่: ' + caption)
        set_thai(r, size=13, color=(107,114,128))

def divider(doc, color='1E40AF'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(doc, headers, rows, col_widths=None, header_color='1E40AF'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, txt in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(txt)
        set_thai(r, size=15, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, header_color)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            set_thai(r, size=15)
            if ri % 2 == 0:
                shade_cell(cell, 'EFF6FF')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

# ─── cover ───────────────────────────────────────────────────────────────────

def cover(doc):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('รายงานการพัฒนาระบบสารสนเทศ'), size=28, bold=True, color=(30,64,175))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('ระบบติดตามแผนงานโครงการและงบประมาณ'), size=24, bold=True, color=(30,64,175))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('Project Plan & Budget Tracking System'), size=18, color=(107,114,128))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('สำนักวิทยบริการ มหาวิทยาลัยนครพนม'), size=18, bold=True, color=(55,65,81))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('Academic Resource Center, Nakhon Phanom University'), size=15, color=(107,114,128))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('จัดทำโดย: งานเทคโนโลยีสารสนเทศ'), size=15, color=(55,65,81))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai(p.add_run('เวอร์ชัน 1.0  |  ปีงบประมาณ 2568'), size=15, color=(107,114,128))

# ─── ส่วนที่ 1: สรุปผู้บริหาร ─────────────────────────────────────────────────

def section1(doc):
    h(doc, 'ส่วนที่ 1  สรุปสำหรับผู้บริหาร', level=1)
    divider(doc)

    # 1.1
    h(doc, '1.1  ระบบนี้คืออะไร', level=2)
    para(doc,
        'ระบบติดตามแผนงานโครงการและงบประมาณ (Project Plan & Budget Tracking System) '
        'เป็นระบบเว็บแอปพลิเคชันที่พัฒนาขึ้นโดยงานเทคโนโลยีสารสนเทศ สำนักวิทยบริการ '
        'มหาวิทยาลัยนครพนม เพื่อแก้ไขข้อจำกัดของระบบ ERP เดิมที่จำกัดการเข้าถึงข้อมูล '
        'แผนโครงการเฉพาะเจ้าหน้าที่แผนเท่านั้น ระบบใหม่เปิดให้เจ้าหน้าที่ทุกระดับ '
        'สามารถติดตามโครงการและงบประมาณที่ตนรับผิดชอบได้แบบ Real-time ผ่านเว็บเบราว์เซอร์ '
        'และรับการแจ้งเตือนผ่าน LINE ได้ทันที')

    img(doc, 'sdlc_02_dashboard.png', width_cm=15, caption='หน้า Dashboard ภาพรวมโครงการและงบประมาณ')

    # 1.2
    h(doc, '1.2  ปัญหาที่ระบบแก้ไข', level=2)
    add_table(doc,
        headers=['ปัญหา', 'ผลกระทบ', 'ระดับ'],
        rows=[
            ('เข้าถึงข้อมูลโครงการไม่ได้\n(ERP จำกัดเฉพาะเจ้าหน้าที่แผน)',
             'ต้องสอบถามเจ้าหน้าที่แผนทุกครั้ง\nเสียเวลาและเกิดความล่าช้า', 'สูง'),
            ('พึ่งพาเอกสารกระดาษ (Hardcopy)',
             'เอกสารสูญหาย ข้อมูลคลาดเคลื่อน\nไม่สามารถค้นหาย้อนหลังได้', 'สูง'),
            ('ไม่รู้ยอดงบประมาณคงเหลือรายกิจกรรม',
             'เสี่ยงใช้จ่ายเกินงบโดยไม่ทราบ\nอาจส่งผลต่อแผนงานทั้งโครงการ', 'สูงมาก'),
            ('ไม่มีระบบแจ้งเตือนอัตโนมัติ',
             'งานล่าช้า พลาดกำหนดส่ง\nไม่รู้ว่างบใกล้หมด', 'ปานกลาง'),
        ],
        col_widths=[5.5, 7, 2.5],
    )

    # 1.3
    h(doc, '1.3  สิ่งที่พัฒนา — 6 Module หลัก', level=2)
    modules = [
        ('1', 'Authentication & RBAC', 'ยืนยันตัวตนผ่าน NPU Active Directory + ควบคุมสิทธิ์ 4 ระดับ'),
        ('2', 'Project Management', 'จัดการโครงการ กิจกรรม งบประมาณ พร้อมระบบโอนงบและแนบเอกสาร'),
        ('3', 'Budget & Expense', 'บันทึกรายจ่าย ติดตามงบแบบ Real-time ระบบอนุมัติ/ปฏิเสธ'),
        ('4', 'Reports', 'รายงานงบประมาณ ส่งออก Excel/PDF พร้อมฟอนต์ภาษาไทย'),
        ('5', 'LINE Notification', 'แจ้งเตือนอัตโนมัติผ่าน LINE เมื่องบใกล้หมดหรือใกล้กำหนด'),
        ('6', 'Admin & Audit', 'จัดการผู้ใช้ แผนก ปีงบประมาณ พร้อม Audit Log ทุกการเปลี่ยนแปลง'),
    ]
    add_table(doc,
        headers=['#', 'Module', 'ความสามารถหลัก'],
        rows=modules,
        col_widths=[1, 4.5, 9.5],
    )

    # 1.4
    h(doc, '1.4  ผลลัพธ์ที่วัดได้', level=2)
    add_table(doc,
        headers=['ตัวชี้วัด (KPI)', 'เป้าหมาย', 'ผลจริง'],
        rows=[
            ('เจ้าหน้าที่เข้าถึงข้อมูลโครงการตนเองได้', '100%', '✓ ทุกคนเข้าได้ตลอด 24 ชม.'),
            ('ลดการสอบถามเจ้าหน้าที่แผน', '≥ 70%', '✓ ข้อมูลเปิดให้ดูเองได้ทันที'),
            ('ติดตามงบประมาณคงเหลือแบบ Real-time', '100%', '✓ แสดงยอดจัดสรร/ใช้ไป/คงเหลือ'),
            ('ส่งแจ้งเตือน LINE สำเร็จ', '≥ 98%', '✓ ส่งอัตโนมัติ 4 ประเภท'),
            ('รองรับผู้ใช้งานหลาย Role', '4 ระดับ', '✓ Staff / Planner / Head / Admin'),
            ('ส่งออกรายงาน PDF/Excel', 'ทำได้', '✓ ReportLab + openpyxl พร้อมใช้'),
        ],
        col_widths=[7, 3, 5],
    )

    para(doc, 'ระบบเปิดใช้งานจริงที่: https://lib.npu.ac.th/projects/',
         size=14, color=(30,64,175), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ─── ส่วนที่ 2: SDLC ──────────────────────────────────────────────────────────

def section2(doc):
    page_break(doc)
    h(doc, 'ส่วนที่ 2  กระบวนการพัฒนาระบบ (SDLC)', level=1)
    divider(doc)
    para(doc,
        'การพัฒนาระบบใช้แนวทาง Incremental Development แบ่งออกเป็น 4 Phase '
        'ดำเนินการตั้งแต่เดือนกุมภาพันธ์ 2568 จนแล้วเสร็จและเปิดใช้งานจริง '
        'รวมระยะเวลาการพัฒนาประมาณ 6 สัปดาห์')

    # Timeline table
    h(doc, '2.0  ภาพรวมแผนการดำเนินงาน', level=2)
    add_table(doc,
        headers=['Phase', 'กิจกรรม', 'ระยะเวลา', 'ผลลัพธ์'],
        rows=[
            ('Phase 1\nวิเคราะห์', 'วิเคราะห์ปัญหา\nกำหนด Requirements\nออกแบบ DB & UI', 'สัปดาห์ 1–2', 'ข้อกำหนดระบบ\nER Diagram\nWireframe'),
            ('Phase 2\nออกแบบ & พัฒนา', 'Models & Migrations\nAuth & RBAC\nCRUD ทุก Module\nDashboard', 'สัปดาห์ 2–4', 'ระบบจัดการโครงการ\nBudget Tracking\nExpense Workflow'),
            ('Phase 3\nNotifications & Reports', 'LINE API Integration\nCelery Scheduler\nรายงาน Excel/PDF\nTimeline View', 'สัปดาห์ 4–5', 'LINE แจ้งเตือนอัตโนมัติ\nส่งออกรายงาน\nปฏิทิน Gantt'),
            ('Phase 4\nDeploy & ปรับปรุง', 'UAT & Bug Fix\nDeploy บน IIS/Windows\nปรับปรุง UI/UX\nจัดทำคู่มือ', 'สัปดาห์ 5–6', 'ระบบพร้อมใช้งาน\nURL จริง\nคู่มือผู้ใช้'),
        ],
        col_widths=[3, 5.5, 3, 3.5],
    )

    # Phase 1
    h(doc, '2.1  Phase 1 — วิเคราะห์ความต้องการ (Requirements Analysis)', level=2)

    h(doc, 'ผู้มีส่วนได้ส่วนเสีย (Stakeholders)', level=3)
    add_table(doc,
        headers=['กลุ่ม', 'บทบาท', 'ความต้องการหลัก'],
        rows=[
            ('เจ้าหน้าที่ผู้ปฏิบัติงาน', 'ผู้ใช้งานหลัก', 'ดูโครงการตนเอง บันทึกรายจ่าย รับแจ้งเตือน'),
            ('เจ้าหน้าที่แผน', 'ผู้ดูแลข้อมูล', 'จัดการโครงการทุกโครงการในแผนก'),
            ('หัวหน้างาน', 'ผู้อนุมัติ', 'อนุมัติรายจ่าย ดู Dashboard ภาพรวม'),
            ('ผู้บริหาร', 'ผู้ตัดสินใจ', 'ภาพรวมงบประมาณ สถานะโครงการทั้งหมด'),
            ('ผู้ดูแลระบบ', 'System Admin', 'จัดการผู้ใช้ ตั้งค่าระบบ ดู Audit Log'),
        ],
        col_widths=[4, 3.5, 7.5],
    )

    h(doc, 'ความต้องการหลักของระบบ (Functional Requirements)', level=3)
    reqs = [
        'ระบบต้องรองรับการ Login ผ่าน NPU Active Directory',
        'กำหนดสิทธิ์การเข้าถึงข้อมูล 4 ระดับ ตามหลัก Role-Based Access Control (RBAC)',
        'แสดงยอดงบประมาณจัดสรร ยอดใช้จ่าย และยอดคงเหลือแบบ Real-time',
        'รองรับการบันทึกรายจ่ายพร้อมระบบอนุมัติ 2 ขั้นตอน',
        'ส่งการแจ้งเตือนอัตโนมัติผ่าน LINE เมื่องบประมาณใกล้หมดหรือใกล้กำหนด',
        'สร้างรายงานสรุปงบประมาณ ส่งออกเป็น Excel และ PDF ได้',
        'รองรับการโอนงบประมาณระหว่างกิจกรรม (Budget Transfer)',
        'บันทึก Audit Log ทุกการเปลี่ยนแปลงข้อมูลสำคัญ',
    ]
    for r in reqs:
        bullet(doc, r)

    # Phase 2
    h(doc, '2.2  Phase 2 — ออกแบบระบบ (System Design)', level=2)

    h(doc, 'สถาปัตยกรรมระบบ (3-Tier Architecture)', level=3)
    add_table(doc,
        headers=['Layer', 'เทคโนโลยี', 'หน้าที่'],
        rows=[
            ('Presentation Layer', 'Tailwind CSS 3.x, HTML5, JavaScript', 'แสดงผลหน้าเว็บ Dashboard Forms'),
            ('Application Layer', 'Django 5.1 (Python 3.12)', 'Business Logic, RBAC, API, Signals'),
            ('Data Layer', 'MySQL 8.0 (utf8mb4)', 'จัดเก็บข้อมูลโครงการ งบประมาณ ผู้ใช้'),
            ('Notification Layer', 'LINE Messaging API v2', 'Push Message & Flex Message'),
            ('Task Queue', 'Celery 5.4 + Redis 7', 'Background Tasks, Scheduled Alerts'),
            ('Web Server', 'Waitress + IIS (ARR)', 'Production Server, Reverse Proxy, SSL'),
        ],
        col_widths=[4, 5.5, 5.5],
    )

    h(doc, 'การควบคุมสิทธิ์การเข้าถึง (RBAC)', level=3)
    add_table(doc,
        headers=['บทบาท', 'ขอบเขต', 'ความสามารถพิเศษ'],
        rows=[
            ('เจ้าหน้าที่ (Staff)', 'โครงการตนเองเท่านั้น', 'บันทึกรายจ่าย'),
            ('เจ้าหน้าที่แผน (Planner)', 'ทุกโครงการในแผนก', 'สร้าง/แก้ไขโครงการ, อนุมัติรายจ่าย'),
            ('หัวหน้างาน (Head)', 'ทุกโครงการในแผนก', 'อนุมัติ/ปฏิเสธรายจ่าย, โอนงบ'),
            ('ผู้ดูแลระบบ (Admin)', 'ทุกโครงการทุกแผนก', 'จัดการผู้ใช้ ตั้งค่า ดู Audit Log'),
        ],
        col_widths=[4.5, 5, 5.5],
    )

    h(doc, 'โครงสร้างฐานข้อมูลหลัก', level=3)
    para(doc, 'ฐานข้อมูลออกแบบตาม Relational Database Design ประกอบด้วยตารางหลัก 8 ตาราง:')
    add_table(doc,
        headers=['ตาราง', 'ข้อมูลที่เก็บ', 'ความสัมพันธ์'],
        rows=[
            ('FiscalYear', 'ปีงบประมาณ วันเริ่ม-สิ้นสุด', 'Parent ของ Project'),
            ('Department', 'แผนก รหัสแผนก', 'Parent ของ Project และ UserProfile'),
            ('Project', 'โครงการ รหัส ชื่อ งบรวม สถานะ', 'FK → FiscalYear, Department'),
            ('Activity', 'กิจกรรม งบจัดสรร สถานะ ผู้รับผิดชอบ', 'FK → Project'),
            ('Expense', 'รายจ่าย จำนวน สถานะอนุมัติ', 'FK → Activity'),
            ('BudgetTransfer', 'การโอนงบระหว่างกิจกรรม', 'FK → Activity (from/to)'),
            ('UserProfile', 'ข้อมูลผู้ใช้ Role LINE UID', 'OneToOne → User'),
            ('LINENotificationLog', 'ประวัติการแจ้งเตือน', 'FK → User, Project, Activity'),
        ],
        col_widths=[3.5, 6, 5.5],
    )

    # Phase 3
    h(doc, '2.3  Phase 3 — การพัฒนาระบบ (Implementation)', level=2)

    h(doc, 'เทคโนโลยีที่ใช้และเหตุผล', level=3)
    add_table(doc,
        headers=['เทคโนโลยี', 'เวอร์ชัน', 'เหตุผลที่เลือก'],
        rows=[
            ('Django Framework', '5.1', 'มีระบบ Auth, ORM, Admin ในตัว ใช้อยู่แล้ว'),
            ('MySQL', '8.0', 'เสถียร รองรับ Transaction ใช้อยู่แล้ว'),
            ('Tailwind CSS', '3.x (CDN)', 'ยืดหยุ่น Responsive ไม่ต้อง Build'),
            ('Celery + Redis', '5.4 + 7', 'Background Tasks แจ้งเตือนอัตโนมัติ'),
            ('LINE Messaging API', 'v2', 'มี License อยู่แล้ว ทุกคนใช้ LINE'),
            ('ReportLab', '4.4', 'PDF ภาษาไทย ฟอนต์ TH Sarabun New'),
            ('openpyxl', '3.x', 'ส่งออก Excel แบบ Styled'),
            ('Waitress + IIS', '-', 'Production บน Windows Server'),
        ],
        col_widths=[4, 2.5, 8.5],
    )

    h(doc, 'หน้าจอที่พัฒนา', level=3)
    para(doc, 'ระบบประกอบด้วยหน้าจอหลักดังนี้:')
    screens = [
        'Dashboard — ภาพรวมโครงการ งบประมาณ รายการที่ต้องติดตาม',
        'รายการโครงการ — ค้นหา กรอง แสดงกิจกรรมย่อยแบบ Expandable',
        'รายละเอียดโครงการ — ข้อมูลครบ พร้อม Progress Bar แยกแหล่งเงิน',
        'รายละเอียดกิจกรรม — บันทึกรายจ่าย แนบไฟล์ ดู PDF/รูปภาพใน Popup',
        'ปฏิทินโครงการ — Gantt Chart แสดงทุกโครงการตลอดปีงบประมาณ',
        'ภาพรวมผู้บริหาร — Dashboard สรุประดับผู้บริหาร',
        'งานของฉัน — รายการงานที่ต้องดำเนินการ กิจกรรมเกินกำหนด',
        'รายงาน — งบประมาณ / รายจ่าย พร้อมส่งออก Excel/PDF',
        'จัดการผู้ใช้ — CRUD ผู้ใช้ กำหนด Role เชื่อม LINE',
    ]
    for s in screens:
        bullet(doc, s)

    img(doc, 'sdlc_03_project_list.png', width_cm=15, caption='หน้ารายการโครงการพร้อมกิจกรรมย่อย')
    img(doc, 'sdlc_05_activity_budget.png', width_cm=15, caption='หน้ากิจกรรม — ติดตามงบประมาณแบบ Real-time')

    h(doc, 'ระบบแจ้งเตือน LINE', level=3)
    add_table(doc,
        headers=['ประเภทแจ้งเตือน', 'เงื่อนไขการส่ง', 'รูปแบบ'],
        rows=[
            ('แจ้งเตือนงบประมาณ', 'ใช้งบถึง Threshold ที่กำหนด (ค่าเริ่มต้น 80%)', 'Flex Message'),
            ('แจ้งเตือนกำหนดการ', '7 วัน และ 3 วัน ก่อนถึงกำหนด', 'Text Message'),
            ('แจ้งเตือนสถานะ', 'เมื่อโครงการ/กิจกรรมเปลี่ยนสถานะ', 'Text Message'),
            ('แจ้งเตือนอนุมัติ', 'เมื่อรายจ่ายได้รับการอนุมัติหรือปฏิเสธ', 'Text Message'),
        ],
        col_widths=[4.5, 6.5, 4],
    )

    # Phase 4
    h(doc, '2.4  Phase 4 — ทดสอบและติดตั้งใช้งาน (Testing & Deployment)', level=2)

    h(doc, 'การทดสอบระบบ', level=3)
    add_table(doc,
        headers=['ระดับการทดสอบ', 'วิธีการ', 'ผล'],
        rows=[
            ('Functional Testing', 'ทดสอบ CRUD ทุก Module ตรวจสอบ RBAC', 'ผ่าน'),
            ('Budget Logic Testing', 'ทดสอบการคำนวณงบ การโอนงบ การเกินงบ', 'ผ่าน'),
            ('LINE Notification Testing', 'ทดสอบการส่งทุก 4 ประเภท', 'ผ่าน'),
            ('UAT (User Acceptance)', 'เจ้าหน้าที่ใช้งานจริงกับข้อมูลจริง', 'ผ่าน'),
            ('Security Testing', 'ทดสอบ RBAC, CSRF, SQL Injection', 'ผ่าน'),
        ],
        col_widths=[5, 7, 3],
    )

    h(doc, 'การติดตั้งระบบ Production', level=3)
    add_table(doc,
        headers=['รายการ', 'รายละเอียด'],
        rows=[
            ('URL ที่ใช้งาน', 'https://lib.npu.ac.th/projects/'),
            ('Server', 'Windows Server (IIS + ARR + Waitress)'),
            ('SSL Certificate', '*.npu.ac.th — DigiCert (ถึงกุมภาพันธ์ 2570)'),
            ('Database', 'MySQL 8.0 — 202.29.55.213'),
            ('Service Management', 'NSSM (Non-Sucking Service Manager)'),
            ('Authentication', 'NPU Active Directory + Local fallback'),
        ],
        col_widths=[5, 10],
    )

    img(doc, 'sdlc_07_timeline.png', width_cm=15, caption='ปฏิทิน Gantt — แสดงแผนงานทุกโครงการตลอดปีงบประมาณ')
    img(doc, 'sdlc_06_report.png', width_cm=15, caption='หน้ารายงานงบประมาณ พร้อมส่งออก Excel/PDF')

# ─── ภาคผนวก ──────────────────────────────────────────────────────────────────

def appendix(doc):
    page_break(doc)
    h(doc, 'ภาคผนวก — รายละเอียดเพิ่มเติม', level=1)
    divider(doc)

    h(doc, 'ภาคผนวก ก  โครงสร้างแอปพลิเคชัน', level=2)
    add_table(doc,
        headers=['App', 'Module', 'ฟังก์ชันหลัก'],
        rows=[
            ('accounts', 'Auth, UserProfile, Department', 'Login NPU AD, RBAC, จัดการผู้ใช้'),
            ('projects', 'FiscalYear, Project, Activity', 'CRUD โครงการและกิจกรรม, Timeline'),
            ('budget', 'Expense, BudgetTransfer', 'บันทึกรายจ่าย, อนุมัติ, โอนงบ'),
            ('notifications', 'LINEService, NotificationLog', 'Push LINE, Celery Tasks'),
            ('reports', 'Budget/Expense Reports', 'สร้าง HTML/Excel/PDF'),
            ('dashboard', 'Dashboard, Executive View', 'ภาพรวมตาม Role'),
        ],
        col_widths=[3, 5, 7],
    )

    h(doc, 'ภาคผนวก ข  URL หลักของระบบ', level=2)
    add_table(doc,
        headers=['URL', 'หน้าจอ', 'สิทธิ์'],
        rows=[
            ('/projects/', 'Dashboard', 'ทุก Role'),
            ('/projects/projects/', 'รายการโครงการ', 'ทุก Role'),
            ('/projects/projects/<pk>/', 'รายละเอียดโครงการ', 'ทุก Role'),
            ('/projects/projects/timeline/', 'ปฏิทิน Gantt', 'ทุก Role'),
            ('/projects/budget/', 'รายการรายจ่าย', 'ทุก Role'),
            ('/projects/budget/approvals/', 'อนุมัติรายจ่าย', 'Head, Admin'),
            ('/projects/executive/', 'ภาพรวมผู้บริหาร', 'Head, Admin'),
            ('/projects/reports/budget/', 'รายงานงบประมาณ', 'Planner, Head, Admin'),
            ('/projects/my-tasks/', 'งานของฉัน', 'ทุก Role'),
            ('/projects/accounts/manage/', 'จัดการระบบ', 'Admin'),
        ],
        col_widths=[6, 5, 4],
    )

    h(doc, 'ภาคผนวก ค  Screenshots ระบบ', level=2)
    img(doc, 'sdlc_01_login.png', width_cm=13, caption='หน้าเข้าสู่ระบบ (Login)')
    img(doc, 'sdlc_04_project_detail.png', width_cm=13, caption='รายละเอียดโครงการ')
    img(doc, 'sdlc_08_executive.png', width_cm=13, caption='ภาพรวมผู้บริหาร')
    img(doc, 'sdlc_09_admin.png', width_cm=13, caption='หน้าผู้ดูแลระบบ')
    img(doc, 'sdlc_10_profile_line.png', width_cm=13, caption='ตั้งค่าการแจ้งเตือน LINE')

# ─── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_page_a4(doc)

    # default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style.font.size = Pt(15)

    cover(doc)
    page_break(doc)
    section1(doc)
    section2(doc)
    appendix(doc)

    doc.save(OUTPUT)
    print(f'✓ บันทึกแล้ว: {OUTPUT}')

if __name__ == '__main__':
    main()
